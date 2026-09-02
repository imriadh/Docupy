import io
import json
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import markdown
import matplotlib
matplotlib.use("Agg")  # Non-gui backend
import matplotlib.pyplot as plt
import nbformat as nbf
from xhtml2pdf import pisa

PYTHON_LANGUAGES = {"python", "py"}


# --- PARSERS & UTILITIES ---


def fix_math_equations(text: str) -> str:
    """Converts LaTeX delimiters to standard Markdown math formatting."""
    text = re.sub(r"\\\[(.*?)\\\]", r"$$\1$$", text, flags=re.DOTALL)
    text = re.sub(r"\\\((.*?)\\\)", r"$\1$", text, flags=re.DOTALL)
    return text


def parse_ai_text_blocks(raw_text: str):
    """Splits raw text into alternating markdown and code metadata chunks."""
    cleaned_text = fix_math_equations(raw_text)
    pattern = r"```([a-zA-Z0-9_\-\+]*)\n(.*?)```"

    blocks = []
    last_end = 0

    for match in re.finditer(pattern, cleaned_text, re.DOTALL):
        start, end = match.span()
        lang = match.group(1).strip().lower()
        content = match.group(2).strip()

        markdown_text = cleaned_text[last_end:start].strip()
        if markdown_text:
            blocks.append({"type": "markdown", "content": markdown_text})

        blocks.append({"type": "code", "lang": lang, "content": content})
        last_end = end

    remaining_text = cleaned_text[last_end:].strip()
    if remaining_text:
        blocks.append({"type": "markdown", "content": remaining_text})

    return blocks


# --- MATH TO SVG / OMML CONVERTERS ---


def latex_to_png_bytes(formula: str, fontsize=14, is_block=False) -> io.BytesIO:
    """Renders LaTeX formula to a PNG image buffer for Word insertion."""
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0.0)

    # Render latex expression
    text = fig.text(
        0, 0, f"${formula}$", fontsize=fontsize, color="black", usetex=False
    )

    buffer = io.BytesIO()
    fig.savefig(
        buffer,
        format="png",
        bbox_inches="tight",
        pad_inches=0.05,
        transparent=True,
        dpi=300,
    )
    plt.close(fig)
    buffer.seek(0)
    return buffer


def replace_math_with_svg_for_pdf(text: str) -> str:
    """Replaces $$math$$ and $math$ with inline SVG tags for xhtml2pdf rendering."""

    def render_block_svg(match):
        formula = match.group(1).strip()
        img_buffer = latex_to_png_bytes(formula, fontsize=16, is_block=True)
        import base64

        b64 = base64.b64encode(img_buffer.getvalue()).decode("utf-8")
        return f'<div style="text-align:center; margin: 10px 0;"><img src="data:image/png;base64,{b64}" /></div>'

    def render_inline_svg(match):
        formula = match.group(1).strip()
        img_buffer = latex_to_png_bytes(formula, fontsize=12, is_block=False)
        import base64

        b64 = base64.b64encode(img_buffer.getvalue()).decode("utf-8")
        return f'<img src="data:image/png;base64,{b64}" style="vertical-align:middle;" />'

    # Convert block equations
    text = re.sub(r"\$\$(.*?)\$\$", render_block_svg, text, flags=re.DOTALL)
    # Convert inline equations
    text = re.sub(r"\$(.*?)\$", render_inline_svg, text)
    return text


# --- CORE CONVERTERS ---


def txt_to_ipynb(raw_text: str) -> tuple[str, int]:
    """Converts AI raw text into an IPYNB string and returns cell count."""
    blocks = parse_ai_text_blocks(raw_text)
    nb = nbf.v4.new_notebook()
    cells = []

    for block in blocks:
        if block["type"] == "markdown":
            cells.append(nbf.v4.new_markdown_cell(block["content"]))
        elif block["type"] == "code":
            lang = block["lang"]
            content = block["content"]
            if lang in PYTHON_LANGUAGES:
                cells.append(nbf.v4.new_code_cell(content))
            else:
                lang_str = f"```{lang}\n{content}\n```"
                cells.append(nbf.v4.new_markdown_cell(lang_str))

    nb["cells"] = cells
    return nbf.writes(nb), len(cells)


def ipynb_to_md(nb_json: dict, include_outputs: bool = False) -> str:
    """Converts Notebook JSON data to Markdown format."""
    md_content = []

    for cell in nb_json.get("cells", []):
        cell_type = cell.get("cell_type")
        source = (
            "".join(cell.get("source", []))
            if isinstance(cell.get("source"), list)
            else cell.get("source", "")
        )

        if cell_type == "markdown":
            md_content.append(f"{source}\n\n")
        elif cell_type == "code":
            md_content.append(f"```python\n{source}\n```\n\n")

            if include_outputs and "outputs" in cell:
                for output in cell["outputs"]:
                    if output.get("output_type") == "stream":
                        out_text = "".join(output.get("text", []))
                        md_content.append(
                            f"> **Output:**\n```text\n{out_text}\n```\n\n"
                        )

    return "".join(md_content)


def text_to_docx(raw_text: str) -> io.BytesIO:
    """Converts AI text with Math equations and Code blocks into styled Word (.docx)."""
    blocks = parse_ai_text_blocks(raw_text)
    doc = Document()

    for block in blocks:
        if block["type"] == "markdown":
            paragraphs = block["content"].split("\n\n")
            for p_text in paragraphs:
                p_text_stripped = p_text.strip()
                if not p_text_stripped:
                    continue

                # Heading detection
                if p_text_stripped.startswith("#"):
                    level = min(p_text_stripped.count("#", 0, 4), 3)
                    heading_text = p_text_stripped.lstrip("#").strip()
                    doc.add_heading(heading_text, level=level)
                else:
                    p = doc.add_paragraph()

                    # Handle block math ($$ ... $$)
                    if p_text_stripped.startswith("$$") and p_text_stripped.endswith("$$"):
                        formula = p_text_stripped.strip("$").strip()
                        img_buf = latex_to_png_bytes(formula, fontsize=14, is_block=True)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(img_buf)
                    else:
                        # Parse inline math ($ ... $) and text
                        parts = re.split(r"(\$.*?\$)", p_text_stripped)
                        for part in parts:
                            if part.startswith("$") and part.endswith("$"):
                                formula = part.strip("$").strip()
                                img_buf = latex_to_png_bytes(
                                    formula, fontsize=11, is_block=False
                                )
                                p.add_run().add_picture(img_buf)
                            else:
                                p.add_run(part)

        elif block["type"] == "code":
            # Render Code Blocks cleanly with shaded table cell
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False

            # Shading XML for grey background
            cell = table.cell(0, 0)
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:val"), "clear")
            shading_elm.set(qn("w:color"), "auto")
            shading_elm.set(qn("w:fill"), "F1F5F9")
            cell._tc.get_or_add_tcPr().append(shading_elm)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            run = p.add_run(block["content"])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 41, 59)

            doc.add_paragraph()  # spacing after block

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def text_to_pdf(raw_text: str) -> bytes:
    """Renders text with rendered LaTeX Math and syntax-highlighted code to PDF."""
    cleaned_text = fix_math_equations(raw_text)

    # Convert math expressions to embedded image tags first
    text_with_math_images = replace_math_with_svg_for_pdf(cleaned_text)

    # Convert Markdown to HTML
    html_body = markdown.markdown(
        text_with_math_images, extensions=["fenced_code", "tables"]
    )

    # ReportLab treats backslashes as markup escapes, so literal LaTeX-style
    # sequences like \le must be doubled before XHTML-to-PDF rendering.
    html_body = html_body.replace("\\", "\\\\")

    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {{
                size: a4 portrait;
                margin: 18mm 12mm;
            }}
            body {{
                font-family: Helvetica, Arial, sans-serif;
                font-size: 10pt;
                line-height: 1.6;
                color: #1e293b;
            }}
            h1, h2, h3 {{ 
                color: #0f172a; 
                margin-top: 1.2em;
            }}
            pre {{
                background-color: #f8fafc;
                border: 1px solid #e2e8f0;
                padding: 10px;
                font-family: Courier, monospace;
                font-size: 9pt;
                white-space: pre-wrap;
                word-wrap: break-word;
            }}
            code {{
                background-color: #f1f5f9;
                padding: 2px 4px;
                font-family: Courier, monospace;
                font-size: 9pt;
            }}
        </style>
    </head>
    <body>
        {html_body}
    </body>
    </html>
    """

    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=pdf_buffer)

    if pisa_status.err:
        raise RuntimeError("PDF generation failed.")

    return pdf_buffer.getvalue()
